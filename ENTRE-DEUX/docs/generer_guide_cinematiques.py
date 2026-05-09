# ─────────────────────────────────────────────────────────────────────────────
#  LIMINAL (ENTRE-DEUX) — Générateur de guide Word pour les cinématiques
# ─────────────────────────────────────────────────────────────────────────────
#
#  Lancer :  py docs/generer_guide_cinematiques.py
#  Produit : docs/Guide_Cinematiques_LIMINAL.docx
#
#  Pourquoi un script et pas juste un .docx commit ?
#    Pour pouvoir REGÉNÉRER le doc quand on ajoute une fonctionnalité
#    (nouvelle action de cinématique, nouveau type d'event, etc.). Le
#    contenu est en clair dans ce script, donc lisible/modifiable même
#    sans Word.
#
#  Dépendance : python-docx ( py -m pip install python-docx )
# ─────────────────────────────────────────────────────────────────────────────

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


VIOLET       = RGBColor(0x6E, 0x5A, 0xC8)
VIOLET_FONCE = RGBColor(0x40, 0x30, 0x80)
DORE         = RGBColor(0xC8, 0xA8, 0x40)
GRIS         = RGBColor(0x66, 0x66, 0x77)


def h1(doc, txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs:
        r.font.color.rgb = VIOLET_FONCE


def h2(doc, txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs:
        r.font.color.rgb = VIOLET


def h3(doc, txt):
    p = doc.add_heading(txt, level=3)
    for r in p.runs:
        r.font.color.rgb = VIOLET


def para(doc, txt, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    return p


def code(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x30)
    # Fond gris clair via shading. python-docx ne le fait pas trivialement,
    # on s'en passe — Consolas + indentation suffisent visuellement.
    p.paragraph_format.left_indent = Inches(0.3)
    return p


def bullet(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")


def numbered(doc, txt):
    return doc.add_paragraph(txt, style="List Number")


def hr(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.color.rgb = GRIS


def kvtable(doc, header, rows, widths=None):
    """Petite table 2-3 colonnes avec en-tête violet."""
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Light Grid Accent 4"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = VIOLET_FONCE
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


# ═════════════════════════════════════════════════════════════════════════════
#  CONSTRUCTION DU DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # ── Marges raisonnables ────────────────────────────────────────────────
    for s in doc.sections:
        s.top_margin    = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin   = Inches(0.9)
        s.right_margin  = Inches(0.9)

    # ── Page de titre ──────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LIMINAL (ENTRE-DEUX)")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = VIOLET_FONCE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Guide complet des cinématiques et des PNJ")
    r.italic = True
    r.font.size = Pt(16)
    r.font.color.rgb = VIOLET

    doc.add_paragraph()
    para(doc, ("Ce guide explique comment créer des scènes interactives "
               "dans LIMINAL : faire apparaître un personnage, le faire "
               "marcher, débloquer une compétence, conditionner un dialogue "
               "à un événement passé, etc. Il est destiné à toute "
               "l'équipe — pas besoin de coder en Python pour suivre."))

    para(doc, ("Toutes les actions décrites se font depuis l'ÉDITEUR INTÉGRÉ "
               "AU JEU, avec la souris et le clavier. Un PNJ ne peut plus "
               "n'avoir qu'un seul dialogue par map : il peut en avoir "
               "plusieurs, déblocables au fil de l'histoire."))
    hr(doc)

    # ── Sommaire (manuel, simple liste) ────────────────────────────────────
    h2(doc, "Sommaire")
    bullet(doc, "1. Concepts de base (PNJ, cinématique, story flag)")
    bullet(doc, "2. Ouvrir et naviguer dans les éditeurs")
    bullet(doc, "3. Catalogue des actions de cinématique")
    bullet(doc, "4. Story flags et dialogues conditionnels")
    bullet(doc, "5. Astuces de saisie (presse-papiers, picker, …)")
    bullet(doc, "6. Recettes complètes (exemples pas-à-pas)")
    bullet(doc, "7. Bonus : événements de fin de dialogue (PNJ)")
    bullet(doc, "8. Dépannage (quand ça ne marche pas)")
    doc.add_page_break()

    # ── 1. CONCEPTS ──────────────────────────────────────────────────────
    h1(doc, "1. Concepts de base")

    h2(doc, "PNJ (Personnage Non-Joueur)")
    para(doc, ("Un PNJ est un personnage placé dans une carte qui peut "
               "parler au joueur quand celui-ci s'approche et appuie "
               "sur [E]. Il a :"))
    bullet(doc, "Un nom (affiché flottant au-dessus de lui).")
    bullet(doc, "Un sprite (image fixe ou dossier d'animation).")
    bullet(doc, "Une LISTE de conversations (pas une seule).")
    bullet(doc, "Un mode de dialogue : « boucle_dernier » (par défaut) "
                "ou « restart ».")
    bullet(doc, "Optionnellement : des conditions par conversation "
                "(story flags) et des événements en fin de conversation.")

    para(doc, "Modes de dialogue :", bold=True)
    bullet(doc, "boucle_dernier : on enchaîne les conversations 1, 2, 3… "
                "et une fois la dernière jouée, le PNJ ne dit plus que sa "
                "dernière réplique en boucle. C'est le défaut.")
    bullet(doc, "restart : on cycle 1 → 2 → 3 → 1 → 2 …  Utile pour les "
                "PNJ qui répètent les mêmes choses indéfiniment.")

    h2(doc, "Cinématique")
    para(doc, ("Une cinématique est une suite d'ÉTAPES jouées dans "
               "l'ordre. Le joueur n'a en général pas le contrôle "
               "pendant une cinématique (sauf cas spécial, voir l'action "
               "« Rendre la main au joueur »). Chaque étape est d'un "
               "type donné : afficher un dialogue, faire bouger la "
               "caméra, faire apparaître un PNJ, donner un objet, "
               "attendre N secondes, etc."))
    para(doc, ("Une cinématique se déclenche quand le joueur entre dans "
               "une zone-déclencheur (créée en mode 12 de l'éditeur). "
               "Elle peut aussi être lancée à la fin d'un dialogue PNJ, "
               "ou par n'importe quel autre script."))

    h2(doc, "Story flag")
    para(doc, ("Un story flag est une simple variable booléenne (oui/non) "
               "globale à la partie. Exemples :"))
    bullet(doc, "parchemins_lus = oui  (le joueur a lu les deux parchemins)")
    bullet(doc, "boss_foret_battu = oui  (il a vaincu le boss de la forêt)")
    bullet(doc, "anna_est_amie = oui  (il a aidé Anna)")
    para(doc, ("Les flags sont sauvegardés dans le slot avec le reste de "
               "la progression. On les pose depuis :"))
    bullet(doc, "Une cinématique (action « Poser un story flag »).")
    bullet(doc, "Un événement de fin de dialogue PNJ (event de type « flag »).")
    para(doc, "On les LIT pour conditionner un dialogue PNJ.")
    doc.add_page_break()

    # ── 2. ÉDITEURS ──────────────────────────────────────────────────────
    h1(doc, "2. Ouvrir et naviguer dans les éditeurs")

    h2(doc, "Entrer en mode éditeur")
    numbered(doc, "Lance le jeu (main.py).")
    numbered(doc, "Dans le menu titre, choisis « Mode éditeur ».")
    numbered(doc, "Sélectionne une carte existante ou crée-en une nouvelle.")

    h2(doc, "Éditeur de cinématiques (touche F2)")
    para(doc, ("Une fois en mode éditeur (n'importe quel sous-mode), "
               "appuie sur F2. L'éditeur de cinématiques s'ouvre par-"
               "dessus le jeu."))
    para(doc, "Raccourcis dans l'éditeur de cinématiques :", bold=True)
    bullet(doc, "↑ / ↓  : naviguer dans la liste d'étapes.")
    bullet(doc, "A ou +  : AJOUTER une étape (ouvre le sélecteur de type).")
    bullet(doc, "D ou Suppr ou -  : SUPPRIMER l'étape sélectionnée.")
    bullet(doc, "Entrée  : ÉDITER l'étape sélectionnée (parcourt ses champs).")
    bullet(doc, "Maj+↑ / Maj+↓  : RÉORDONNER (déplacer l'étape).")
    bullet(doc, "Ctrl+S  : SAUVEGARDER la cinématique courante.")
    bullet(doc, "Ctrl+N  : NOUVELLE cinématique (demande un nom).")
    bullet(doc, "Ctrl+O  : OUVRIR une cinématique existante.")
    bullet(doc, "T  : TESTER la cinématique (la joue tout de suite).")
    bullet(doc, "Ctrl+R  : reset le compteur « déjà joué » de cette "
                "cinématique (debug : la rend rejouable).")
    bullet(doc, "Maj+Ctrl+R  : reset TOUS les compteurs.")
    bullet(doc, "Esc  : fermer l'éditeur (retour à l'éditeur de niveau).")

    para(doc, "Pendant l'édition d'un champ :", bold=True)
    bullet(doc, "Entrée  : valider et passer au champ suivant.")
    bullet(doc, "Esc  : annuler (l'étape revient à son état précédent).")
    bullet(doc, "P  : picker — pour les champs X et Y, prend la "
                "position MONDE de la souris à l'écran.")
    bullet(doc, "Ctrl+V  : COLLER depuis le presse-papiers (super utile "
                "pour coller des répliques copiées d'un Word).")
    bullet(doc, "Ctrl+C  : copier le contenu actuel du champ.")

    h2(doc, "Éditeur PNJ (touche F3 sur un PNJ)")
    para(doc, ("Va en mode 11 (PNJ) dans l'éditeur, survole un PNJ, "
               "appuie sur F3. L'éditeur PNJ s'ouvre par-dessus."))
    para(doc, "Niveau « Liste des conversations » :", bold=True)
    bullet(doc, "↑ / ↓  : naviguer.")
    bullet(doc, "Entrée  : entrer dans la conv pour éditer ses lignes.")
    bullet(doc, "A  : AJOUTER une conversation (vide, à remplir).")
    bullet(doc, "D ou Suppr  : SUPPRIMER la conversation sélectionnée.")
    bullet(doc, "W  : basculer le mode (boucle_dernier / restart).")
    bullet(doc, "B  : transformer ce PNJ en SAVE POINT (banc de sauvegarde).")
    bullet(doc, "F  : éditer la CONDITION de la conv sélectionnée "
                "(voir section 4).")
    bullet(doc, "Esc  : fermer.")

    para(doc, "Niveau « Lignes d'une conversation » :", bold=True)
    bullet(doc, "↑ / ↓  : naviguer dans les lignes.")
    bullet(doc, "Entrée  : éditer le TEXTE de la ligne.")
    bullet(doc, "O  : éditer l'ORATEUR (qui parle).")
    bullet(doc, "A  : ajouter une ligne après la sélection.")
    bullet(doc, "D ou Suppr  : supprimer la ligne.")
    bullet(doc, "Maj+↑ / Maj+↓  : réordonner les lignes.")
    bullet(doc, "Esc  : retour à la liste des conversations.")

    para(doc, "Pendant la saisie du TEXTE ou de l'ORATEUR :", bold=True)
    bullet(doc, "Ctrl+V  : coller depuis le presse-papiers (les retours "
                "à la ligne sont remplacés par des espaces — 1 réplique "
                "= 1 ligne).")
    bullet(doc, "Ctrl+C  : copier le contenu.")
    doc.add_page_break()

    # ── 3. CATALOGUE DES ACTIONS ──────────────────────────────────────────
    h1(doc, "3. Catalogue des actions de cinématique")
    para(doc, ("Toutes les actions disponibles dans le sélecteur de type "
               "(touche A dans l'éditeur F2). Elles sont rangées par "
               "famille pour s'y retrouver."))

    h2(doc, "Famille « Temps »")
    kvtable(doc,
        ["Action",     "Champs",                                    "Effet"],
        [
            ["Attendre", "duration (s)", "Pause passive de N secondes."],
            ["Fondu",    "direction (out/in), duration",
                          "Fondu noir vers/depuis. À enchaîner avec un "
                          "changement de carte ou un set_player_pos."],
        ])

    h2(doc, "Famille « Dialogue »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Dialogue", "lignes_texte",
                          "Affiche la boîte de dialogue. L'étape se "
                          "termine quand le joueur a fait défiler "
                          "toutes les lignes."],
        ])
    para(doc, "Format de lignes_texte :", bold=True)
    code(doc, "Bonjour|Anna // Comment vas-tu ?|Anna // Très bien|Théa")
    para(doc, ("Chaque réplique est séparée par //. Avant le | c'est le "
               "TEXTE, après le | c'est l'AUTEUR (qui parle). Si pas "
               "de | l'auteur est vide."))

    h2(doc, "Famille « Caméra »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Caméra → position",   "x, y, duration?, speed?",
                "Lisse la caméra vers (x,y). Si duration omis, reste "
                "focalisée jusqu'à un Caméra → libérer."],
            ["Caméra → PNJ",        "nom, duration?, speed?, follow?",
                "Cible un PNJ. Si follow=1, la caméra suit le PNJ "
                "même s'il bouge ensuite (utile combiné à PNJ marche)."],
            ["Caméra → libérer",    "—",
                "Rend la caméra au joueur (suivi normal)."],
        ])

    h2(doc, "Famille « Personnages »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Joueur marche vers...",  "x, y, speed",
                "Le joueur se déplace automatiquement vers (x,y)."],
            ["PNJ marche vers...",     "nom_pnj, x, y, speed",
                "Le PNJ se déplace vers (x,y). Sa direction "
                "(facing left/right) s'adapte automatiquement, et son "
                "animation passe en walk si le sprite a un dossier walk/."],
            ["Téléporter le joueur",   "x, y",
                "Téléporte instantanément le joueur."],
            ["Faire apparaître un PNJ","nom, x, y, sprite, dialogues, "
                                       "mode, gravité",
                "Crée un PNJ runtime (séraphin qui débarque, etc.)."],
            ["Faire disparaître un PNJ","nom",
                "Retire un PNJ de la scène."],
        ])

    h2(doc, "Famille « Récompenses »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Débloquer une compétence", "value",
                "Pose settings.skill_<value> = True. Valeurs : "
                "double_jump, dash, back_dodge, wall_jump, attack, pogo."],
            ["Donner une luciole", "source",
                "Ajoute 1 compagnon. La source est un identifiant unique "
                "(ex « anna_rite ») pour ne pas redonner si la cinématique "
                "rejoue."],
            ["Donner un item", "name, count",
                "Ajoute count exemplaires de l'item à l'inventaire (avec "
                "stack auto pour Pomme)."],
            ["Donner des pièces", "amount",
                "Ajoute amount pièces au joueur."],
        ])

    h2(doc, "Famille « Story flags »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Poser un story flag", "key, value",
                "Pose game.story_flags[key] = True/False. Lu par les "
                "PNJ pour leurs conditions de dialogue (voir section 4)."],
        ])

    h2(doc, "Famille « Effets »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Secousse",            "amplitude (px), duration (s)",
                "Tremblement de caméra (impact, explosion)."],
            ["Jouer un son",        "nom, volume",
                "Joue un son (déjà chargé via audio/sound_manager)."],
            ["Particules (explosion)", "x, y, nb, couleur",
                "Émet nb particules à (x,y) couleur (r,g,b)."],
        ])

    h2(doc, "Famille « Audio »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Musique : transition", "chemin, volume, fadeout_ms, fadein_ms",
                "Bascule vers une autre piste musicale avec fondu. "
                "Chemin vide = silence (fadeout seul)."],
        ])

    h2(doc, "Famille « Interaction »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Attendre une touche du joueur", "touche, timeout",
                "Pause la cinématique jusqu'à appui. touche = "
                "« any » / « space » / « enter ». timeout=0 → "
                "attend indéfiniment."],
        ])

    h2(doc, "Famille « Hybride cinématique / gameplay »")
    kvtable(doc,
        ["Action", "Champs", "Effet"],
        [
            ["Rendre la main au joueur jusqu'à un point",
             "x, y, radius, timeout",
             "Étape spéciale : le joueur peut SE DÉPLACER, sauter, "
             "attaquer, etc. La cinématique attend qu'il atteigne "
             "(x,y) ± radius. Idéal pour « va ouvrir le tiroir », "
             "« monte à l'échelle ». Le timeout abandonne l'étape "
             "au bout de N secondes (sécurité)."],
        ])
    doc.add_page_break()

    # ── 4. STORY FLAGS / DIALOGUES CONDITIONNELS ───────────────────────
    h1(doc, "4. Story flags et dialogues conditionnels")

    para(doc, ("C'est la fonctionnalité qui rend les PNJ vivants. Avant, "
               "un PNJ disait toujours la même chose. Maintenant, on "
               "peut « débloquer » un nouveau dialogue après un événement."))

    h2(doc, "Étape 1 — Poser un flag dans une cinématique")
    para(doc, ("Dans l'éditeur de cinématique, ajoute l'action « Poser "
               "un story flag » avec :"))
    code(doc, "key   = parchemins_lus\nvalue = 1   (1 = True, 0 = False)")
    para(doc, ("Tu peux poser autant de flags que tu veux. Conventions :"))
    bullet(doc, "Préfère minuscules + underscores : « parchemins_lus », "
                "pas « ParcheminsLus ».")
    bullet(doc, "Préfixe par zone si beaucoup : « foret_boss_battu », "
                "« village_anna_amie ».")
    bullet(doc, "Choisis des noms STABLES — changer le nom invalide "
                "toutes les conditions qui le référencent.")

    h2(doc, "Étape 2 — Conditionner une conversation PNJ")
    para(doc, ("Va sur le PNJ concerné, F3 pour ouvrir l'éditeur, "
               "sélectionne la conversation, appuie sur F. Saisis "
               "l'une des syntaxes suivantes :"))
    kvtable(doc,
        ["Saisie", "Sens"],
        [
            ["(vide)",                 "Toujours dispo (pas de condition)"],
            ["flag:parchemins_lus",    "Dispo si flag True"],
            ["flag:parchemins_lus=0",  "Dispo si flag False ou absent"],
            ["any:k1,k2",              "Dispo si AU MOINS un flag True"],
            ["all:k1,k2,k3",           "Dispo si TOUS les flags sont True"],
        ])

    h2(doc, "Comment le PNJ choisit-il quoi dire ?")
    para(doc, ("Quand le joueur appuie sur [E], le PNJ regarde sa liste "
               "de conversations dans l'ordre, et choisit la PREMIÈRE "
               "qui satisfait sa condition (en partant de l'index "
               "courant). Si aucune ne convient au-delà, il revient à "
               "la dernière dispo trouvée. Concrètement :"))
    bullet(doc, "Conv 1 : pas de condition.")
    bullet(doc, "Conv 2 : flag:parchemins_lus")
    bullet(doc, "Conv 3 : flag:boss_battu")
    para(doc, "Au début de la partie : le PNJ joue la conv 1.")
    para(doc, "Après les parchemins : la conv 2 se débloque ; le PNJ la joue.")
    para(doc, "Après le boss : la conv 3 prend le relais.")
    doc.add_page_break()

    # ── 5. ASTUCES ─────────────────────────────────────────────────────
    h1(doc, "5. Astuces de saisie")

    h2(doc, "Copier-coller depuis ton script (Word / Google Doc)")
    para(doc, ("Tu rédiges ton histoire dans un Word à côté du jeu ? "
               "Tu peux coller des répliques entières dans les éditeurs :"))
    numbered(doc, "Sélectionne la phrase dans Word, fais Ctrl+C.")
    numbered(doc, "Dans l'éditeur du jeu, ouvre un champ texte (édition "
                  "d'une ligne PNJ ou champ d'une cinématique).")
    numbered(doc, "Fais Ctrl+V — le contenu est collé.")
    para(doc, ("Note : dans l'éditeur PNJ, les retours à la ligne d'un "
               "Word multi-lignes sont REMPLACÉS PAR DES ESPACES (1 "
               "réplique = 1 ligne). Pour avoir plusieurs répliques, "
               "il faut soit faire plusieurs collages dans plusieurs "
               "lignes, soit utiliser le champ « lignes_texte » d'une "
               "cinématique avec des // entre répliques."))

    h2(doc, "Picker souris pour les coordonnées")
    para(doc, ("Quand tu édites un champ x ou y dans une cinématique, "
               "tu peux appuyer sur P : la coordonnée MONDE de la "
               "souris est prise automatiquement. Idéal pour positionner "
               "un PNJ où le joueur a cliqué."))

    h2(doc, "Panneau debug des story flags (F5 en mode éditeur)")
    para(doc, ("En mode éditeur, appuie sur F5 pour ouvrir l'overlay "
               "qui liste TOUS les flags posés. Ça permet de :"))
    bullet(doc, "↑ ↓  : naviguer dans les flags.")
    bullet(doc, "T  : basculer le flag sélectionné (True ↔ False) — "
                "très utile pour TESTER les conditions de dialogue PNJ.")
    bullet(doc, "D  : supprimer un flag.")
    bullet(doc, "A  : ajouter un nouveau flag (saisie texte, Ctrl+V "
                "marche).")
    bullet(doc, "F5 ou Esc  : fermer.")
    para(doc, ("Pratique pour ne PAS avoir à rejouer toute une "
               "cinématique pour tester si une conv PNJ se débloque "
               "correctement."), italic=True, color=GRIS)

    h2(doc, "Tester avant de sauver")
    para(doc, ("Touche T dans l'éditeur de cinématique → la cinématique "
               "se joue tout de suite. Si elle a déjà été jouée et que "
               "tu veux la rejouer, fais Ctrl+R d'abord (reset compteur)."))

    h2(doc, "Réordonner sans tout retaper")
    para(doc, ("Maj+↑ / Maj+↓ déplace l'étape ou la ligne sélectionnée "
               "vers le haut ou le bas. Marche dans les deux éditeurs."))
    doc.add_page_break()

    # ── 6. RECETTES ────────────────────────────────────────────────────
    h1(doc, "6. Recettes pas-à-pas")

    h2(doc, "Recette A — Le PNJ qui débloque le double saut")
    para(doc, ("Anna donne le double saut au joueur après lui avoir "
               "parlé. À la 2e fois qu'on lui parle, elle ne dit plus "
               "rien de neuf."))
    numbered(doc, "Place Anna en mode 11.")
    numbered(doc, "F3 sur Anna pour ouvrir l'éditeur PNJ.")
    numbered(doc, "Crée la conversation 1 :  « Tu sembles fatigué, prends "
                  "ce don. »")
    numbered(doc, "Crée la conversation 2 :  « Bonne route, voyageur. »")
    numbered(doc, "Sors au niveau Liste, NE mets PAS de condition (inutile "
                  "ici, mode boucle_dernier suffit).")
    numbered(doc, "Maintenant, on veut que la conv 1 débloque le saut. "
                  "Solution : utiliser un EVENT en fin de conv (voir "
                  "section 7).")
    numbered(doc, "Édite le JSON de la map (alternative simple) ou ajoute "
                  "via une cinématique enchaînée. Pour la suite voir la "
                  "section 7.")

    h2(doc, "Recette B — Le séraphin qui apparaît après les parchemins")
    para(doc, ("Le joueur lit deux parchemins. À la fermeture du second "
               "dialogue, séraphin apparaît derrière lui."))
    para(doc, "Étapes à créer dans l'éditeur de cinématique :", bold=True)
    numbered(doc, "Poser un story flag → key=parchemins_lus, value=1")
    numbered(doc, "Attendre → duration=0.8")
    numbered(doc, "Faire apparaître un PNJ → "
                  "nom=seraphin, x=300, y=200, sprite=seraphin, "
                  "dialogues=« Mais qu'est-ce que tu fais là ?|Séraphin », "
                  "mode=boucle_dernier, gravité=1")
    numbered(doc, "Caméra → PNJ → nom=seraphin, duration=2.0")
    numbered(doc, "Dialogue → lignes_texte=« Mais qu'est-ce que tu fais "
                  "là ?|Séraphin // ...|Toi »")
    numbered(doc, "PNJ marche vers... → nom_pnj=seraphin, x=380, y=200")
    numbered(doc, "Faire disparaître un PNJ → nom=seraphin")
    numbered(doc, "Caméra → libérer")
    numbered(doc, "Fondu → direction=in, duration=0.6")
    para(doc, ("Sauve la cinématique avec Ctrl+S sous un nom comme "
               "« parchemins/seraphin_apparait ». Pour la déclencher, "
               "ajoute une zone-déclencheur (mode 12) qui pointe sur "
               "ce nom."))

    h2(doc, "Recette C — Le coffre qui demande au joueur d'aller chercher "
            "quelque chose")
    para(doc, ("Une cinématique mi-narrative mi-gameplay : on parle au "
               "joueur, on lui dit « va monter à l'échelle », on rend "
               "la main, on attend qu'il y arrive, puis on enchaîne."))
    numbered(doc, "Dialogue → « Tu dois monter là-haut.|Voix »")
    numbered(doc, "Caméra → position → vers le sommet de l'échelle "
                  "(x=820, y=120, duration=1.2)")
    numbered(doc, "Caméra → libérer")
    numbered(doc, "Rendre la main au joueur jusqu'à un point → "
                  "x=820, y=120, radius=40, timeout=60")
    numbered(doc, "Dialogue → « Tu y es. Continue.|Voix »")
    numbered(doc, "Donner une luciole → source=« sommet_echelle »")

    h2(doc, "Recette D — Le PNJ qui change de réplique selon la progression")
    para(doc, ("Un garde au village a 3 attitudes selon ce que le "
               "joueur a fait :"))
    numbered(doc, "F3 sur le garde, crée 3 conversations.")
    numbered(doc, "Sur la conv 1 (par défaut), pas de condition : "
                  "« Va voir, étranger. »")
    numbered(doc, "Sur la conv 2, F → tape « flag:foret_explorée » : "
                  "« On dit que tu as bravé la forêt. »")
    numbered(doc, "Sur la conv 3, F → tape « all:foret_explorée,"
                  "boss_battu » : « Tu es le héros qu'on attendait. »")
    numbered(doc, "Pose les flags via tes cinématiques quand l'événement "
                  "correspondant arrive (action « Poser un story flag »).")
    doc.add_page_break()

    # ── 7. ÉVÉNEMENTS DE FIN DE DIALOGUE ────────────────────────────────
    h1(doc, "7. Événements de fin de dialogue (PNJ)")
    para(doc, ("Plutôt que de passer par une cinématique, un PNJ peut "
               "déclencher un effet automatiquement quand le joueur "
               "ferme une de ses conversations. Pratique pour les "
               "petites récompenses."))

    para(doc, ("Aujourd'hui, ces events s'éditent en JSON directement "
               "(une UI dans l'éditeur PNJ est prévue plus tard). Va "
               "dans le fichier maps/<carte>.json, repère le PNJ par "
               "son nom, et ajoute un champ « events ». C'est une liste "
               "PARALLÈLE aux dialogues : events[i] est déclenché à la "
               "fin de la conv i."))

    para(doc, "Exemple : Anna donne le double saut à la 2e conv.", bold=True)
    code(doc, """{
  "type": "pnj",
  "x": 1200, "y": 400,
  "nom": "Anna",
  "dialogues": [
    [["Tu sembles fatigué...", "Anna"]],
    [["Prends ceci, ça t'aidera.", "Anna"]]
  ],
  "events": [
    [],
    [{"type": "skill", "value": "double_jump"}]
  ]
}""")

    para(doc, "Types d'events supportés :", bold=True)
    kvtable(doc,
        ["Type", "Champs", "Effet"],
        [
            ["skill",   "value",
                "Débloque la compétence settings.skill_<value>."],
            ["luciole", "source",
                "Ajoute 1 luciole. source = identifiant unique."],
            ["coins",   "value",
                "Ajoute value pièces."],
            ["hp",      "value",
                "Soigne (cap = max_hp)."],
            ["max_hp",  "value",
                "Augmente max_hp ET soigne plein."],
            ["item",    "value, count?",
                "Ajoute l'item à l'inventaire (stack auto)."],
            ["flag",    "key, value (true/false)",
                "Pose un story flag global."],
        ])
    doc.add_page_break()

    # ── 8. DÉPANNAGE ─────────────────────────────────────────────────
    h1(doc, "8. Dépannage")

    h2(doc, "« Ma cinématique ne se déclenche pas »")
    bullet(doc, "Vérifie qu'il y a bien une zone-déclencheur (mode 12) "
                "qui pointe sur le nom EXACT de la cinématique.")
    bullet(doc, "Une cinématique ne se rejoue pas par défaut. Pour la "
                "tester à nouveau, appuie sur Ctrl+R dans son éditeur "
                "(reset compteur).")
    bullet(doc, "Si la zone est en dehors du rect du joueur, elle ne se "
                "déclenche pas. Repositionne en mode 12.")

    h2(doc, "« Mon PNJ ne dit pas la nouvelle conversation »")
    bullet(doc, "La condition n'est peut-être pas remplie. Vérifie que "
                "le flag est bien posé (touche T sur la cinématique qui "
                "le pose).")
    bullet(doc, "Si la conv courante ne satisfait pas sa condition, le "
                "PNJ peut « bloquer » sur la précédente. Ordre des "
                "conditions important : la première qui matche gagne, "
                "à partir de l'index courant.")

    h2(doc, "« Le copier-coller (Ctrl+V) ne marche pas »")
    bullet(doc, "Le premier appel peut prendre 1-2 secondes (tkinter "
                "démarre). Réessaie.")
    bullet(doc, "Si rien ne s'affiche, le presse-papiers est peut-être "
                "vide ou contient un format non-texte (image, fichier).")

    h2(doc, "« Le PNJ apparu via npc_spawn n'a pas son sprite »")
    bullet(doc, "Vérifie que le nom du sprite correspond bien à un "
                "fichier ou dossier dans assets/images/pnj/.")
    bullet(doc, "Si tu laisses le champ vide, c'est normal : un "
                "rectangle violet apparaît à la place.")

    h2(doc, "« La caméra reste bloquée après ma cinématique »")
    bullet(doc, "Tu as utilisé Caméra → position ou Caméra → PNJ sans "
                "duration ? Ajoute une étape Caméra → libérer à la fin.")
    doc.add_page_break()

    # ── ANNEXE : Référence rapide ─────────────────────────────────────
    h1(doc, "Annexe — Référence rapide à imprimer")
    h2(doc, "Raccourcis éditeur de cinématique (F2)")
    bullet(doc, "↑↓  navig | A  +  | D / Suppr  -  | Entrée  éditer")
    bullet(doc, "Maj+↑↓  réordonner | Ctrl+S  sauver | Ctrl+O  ouvrir")
    bullet(doc, "T  tester | Ctrl+R  reset compteur | Esc  fermer")

    h2(doc, "Raccourcis éditeur PNJ (F3 sur PNJ)")
    bullet(doc, "↑↓  navig | A  + | D  - | Entrée  éditer | O  orateur")
    bullet(doc, "F  condition | W  mode | B  save point | Esc  fermer")
    bullet(doc, "Maj+↑↓  réordonner les lignes")

    h2(doc, "Champs texte (les deux éditeurs)")
    bullet(doc, "Ctrl+V  coller du presse-papiers")
    bullet(doc, "Ctrl+C  copier le contenu")
    bullet(doc, "P  picker souris (champs x/y des cinématiques)")
    bullet(doc, "Entrée  valider | Esc  annuler")

    h2(doc, "Format des champs spéciaux")
    para(doc, "lignes_texte (action « Dialogue »)", bold=True)
    code(doc, "Texte 1|Auteur1 // Texte 2|Auteur2 // Texte 3|Auteur1")

    para(doc, "dialogues_texte (action « Faire apparaître un PNJ »)", bold=True)
    code(doc, "Conv1 ligne1|A // Conv1 ligne2|B ; Conv2 ligne1|A")
    para(doc, "« ; » entre conversations, « // » entre lignes, « | » "
              "entre texte et auteur.", italic=True, color=GRIS)

    para(doc, "Conditions de dialogue PNJ (touche F)", bold=True)
    code(doc, "(vide)         → toujours dispo\n"
              "flag:k         → dispo si flag True\n"
              "flag:k=0       → dispo si flag False/absent\n"
              "any:k1,k2,k3   → ≥1 flag True\n"
              "all:k1,k2,k3   → tous flags True")

    return doc


# ═════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    doc = build()
    sortie = os.path.join(os.path.dirname(__file__), "Guide_Cinematiques_LIMINAL.docx")
    doc.save(sortie)
    print(f"Document généré : {sortie}")
